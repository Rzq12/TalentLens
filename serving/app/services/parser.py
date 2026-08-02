"""Deterministic document parsing.

This is not an LLM agent and must not become one. PDF/DOCX extraction is a
solved library problem with exact, reproducible character offsets; a model would
paraphrase rather than transcribe, destroying the `(page, start_char, end_char)`
anchoring that the evidence-citation layer depends on.

Failure is loud. A resume is parsed once and reused across every job, so a
guessed or partial parse would silently bias every future score for that
candidate — better to raise and let the document be re-uploaded or OCR'd.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field
from typing import Literal

import fitz
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.config import DOCX_MIME
from app.exceptions import (
    DocumentParseError,
    EmptyDocumentError,
    UnsupportedMediaTypeError,
)
from app.logging import get_logger

logger = get_logger(__name__)

PARSER_VERSION = "pymupdf-1.24+python-docx-1.1/v1"
PAGE_SEPARATOR = "\n\n"
MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER = 20

# Resource ceilings. A file can pass the byte-size check and still be a
# decompression bomb: a few hundred kilobytes of PDF can declare tens of
# thousands of pages, or expand into gigabytes of text. Both are cheap to
# construct and expensive to parse, so the parser bounds its own work rather
# than trusting that upstream size validation was sufficient.
MAX_PAGES = 500
MAX_EXTRACTED_CHARS = 5_000_000  # ~5 MB of text

ParseStatus = Literal["ok", "low_yield", "failed"]


@dataclass(frozen=True, slots=True)
class TextSpan:
    """One styled run of text with the geometry needed to judge visibility.

    The sanitizer cannot detect white-on-white or off-canvas text from a plain
    string, so these attributes are extracted in the same pass as the text.

    Attributes:
        text: The span's characters.
        page: One-based page number.
        size: Font size in points.
        color: Packed sRGB integer as reported by the PDF, or None if unknown.
        bbox: (x0, y0, x1, y1) in page coordinates, or None if unknown.
        page_rect: The page's own (x0, y0, x1, y1), for off-canvas comparison.
    """

    text: str
    page: int
    size: float
    color: int | None
    bbox: tuple[float, float, float, float] | None
    page_rect: tuple[float, float, float, float] | None


@dataclass(frozen=True, slots=True)
class ParsedPage:
    """One page of extracted text, anchored into the whole-document string.

    Attributes:
        page: One-based page number.
        text: The page's extracted text.
        start_char: Inclusive offset of `text` within `ParsedDocument.text`.
        end_char: Exclusive offset of `text` within `ParsedDocument.text`.
    """

    page: int
    text: str
    start_char: int
    end_char: int


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    """The result of parsing one resume or job description.

    Attributes:
        text: Full extracted text, the string all offsets index into.
        pages: Per-page slices with exact offsets.
        page_count: Number of pages found.
        media_type: The media type actually parsed.
        spans: Styled runs with geometry, consumed by the sanitizer.
        metadata: Document metadata, an injection surface in its own right.
        parser_version: Provenance for replay.
        needs_ocr: True when the text layer is too thin to trust.
        parse_status: "ok" when text was recovered, "low_yield" when OCR is needed.
        warnings: Non-fatal notes raised during parsing.
    """

    text: str
    pages: tuple[ParsedPage, ...]
    page_count: int
    media_type: str
    parser_version: str = PARSER_VERSION
    needs_ocr: bool = False
    parse_status: ParseStatus = "ok"
    warnings: tuple[str, ...] = field(default=())
    spans: tuple[TextSpan, ...] = field(default=())
    metadata: dict[str, str] = field(default_factory=dict)


def _assemble(
    pages: list[str],
    media_type: str,
    spans: tuple[TextSpan, ...] = (),
    metadata: dict[str, str] | None = None,
) -> ParsedDocument:
    """Join page texts and compute exact offsets for each page.

    Offsets are derived from the same concatenation that produces `text`, so
    `text[start_char:end_char] == page.text` holds by construction rather than
    by arithmetic that could drift.

    Args:
        pages: Per-page extracted text, in order.
        media_type: The media type that was parsed.
        spans: Styled runs with geometry, when the format exposes them.
        metadata: Document metadata, when the format exposes it.

    Returns:
        The assembled document.
    """
    parts: list[ParsedPage] = []
    cursor = 0
    for index, page_text in enumerate(pages, start=1):
        start = cursor
        end = start + len(page_text)
        parts.append(ParsedPage(page=index, text=page_text, start_char=start, end_char=end))
        cursor = end + len(PAGE_SEPARATOR)

    full_text = PAGE_SEPARATOR.join(pages)
    total_chars = sum(len(p) for p in pages)
    thin = total_chars < MIN_CHARS_PER_PAGE_FOR_TEXT_LAYER * max(len(pages), 1)

    if thin:
        logger.info(
            "parse_low_yield",
            page_count=len(pages),
            total_chars=total_chars,
            needs_ocr=True,
        )

    return ParsedDocument(
        text=full_text,
        pages=tuple(parts),
        page_count=len(pages),
        media_type=media_type,
        needs_ocr=thin,
        parse_status="low_yield" if thin else "ok",
        spans=spans,
        metadata=metadata or {},
    )


def _extract_spans(page: fitz.Page, page_number: int) -> list[TextSpan]:
    """Collect styled runs and their geometry from one PDF page.

    Uses the structured `dict` extraction rather than plain text so the
    sanitizer can see font size, colour, and position — the attributes that
    distinguish text a human can read from text hidden for a model to find.

    Args:
        page: The PyMuPDF page.
        page_number: One-based page number.

    Returns:
        Every non-empty span on the page.
    """
    rect = page.rect
    page_rect = (rect.x0, rect.y0, rect.x1, rect.y1)
    collected: list[TextSpan] = []

    payload = page.get_text("dict")
    for block in payload.get("blocks", ()):
        for line in block.get("lines", ()):
            for span in line.get("spans", ()):
                text = str(span.get("text", ""))
                if not text.strip():
                    continue
                raw_bbox = span.get("bbox")
                bbox: tuple[float, float, float, float] | None = None
                if raw_bbox is not None and len(raw_bbox) == 4:
                    bbox = (
                        float(raw_bbox[0]),
                        float(raw_bbox[1]),
                        float(raw_bbox[2]),
                        float(raw_bbox[3]),
                    )
                collected.append(
                    TextSpan(
                        text=text,
                        page=page_number,
                        size=float(span.get("size", 0.0)),
                        color=span.get("color"),
                        bbox=bbox,
                        page_rect=page_rect,
                    )
                )
    return collected


def _parse_pdf(content: bytes) -> ParsedDocument:
    """Extract text from a PDF using PyMuPDF.

    Args:
        content: Raw PDF bytes.

    Returns:
        The parsed document.

    Raises:
        DocumentParseError: If the PDF cannot be opened or read.
    """
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            if document.page_count > MAX_PAGES:
                logger.warning(
                    "parse_pdf_rejected_page_count",
                    page_count=document.page_count,
                    limit=MAX_PAGES,
                )
                raise DocumentParseError(
                    f"The PDF exceeds the {MAX_PAGES}-page limit."
                )

            pages: list[str] = []
            spans: list[TextSpan] = []
            total_chars = 0
            for index, page in enumerate(document, start=1):
                text = page.get_text("text")
                total_chars += len(text)
                if total_chars > MAX_EXTRACTED_CHARS:
                    logger.warning(
                        "parse_pdf_rejected_text_volume", limit=MAX_EXTRACTED_CHARS
                    )
                    raise DocumentParseError(
                        "The PDF expands to more text than this service will process."
                    )
                pages.append(text)
                spans.extend(_extract_spans(page, index))
            metadata = {
                str(k): str(v) for k, v in (document.metadata or {}).items() if v
            }
    except DocumentParseError:
        raise
    except Exception as err:  # noqa: BLE001 - normalized to a domain error
        logger.warning("parse_pdf_failed", error=str(err), error_type=type(err).__name__)
        raise DocumentParseError("The PDF could not be opened.") from err

    if not pages:
        raise DocumentParseError("The PDF contains no pages.")
    return _assemble(pages, "application/pdf", tuple(spans), metadata)


def _parse_docx(content: bytes) -> ParsedDocument:
    """Extract text from a DOCX using python-docx.

    DOCX has no intrinsic pagination — pagination is a rendering property, not a
    storage one — so the whole document is reported as a single logical page
    rather than inventing page breaks that would make offsets meaningless.

    Args:
        content: Raw DOCX bytes.

    Returns:
        The parsed document.

    Raises:
        DocumentParseError: If the archive is not a readable Word document.
    """
    try:
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs]
    except (
        PackageNotFoundError,
        zipfile.BadZipFile,
        ValueError,
        KeyError,
        OSError,
    ) as err:
        logger.warning("parse_docx_failed", error=str(err), error_type=type(err).__name__)
        raise DocumentParseError("The DOCX could not be opened.") from err

    body = "\n".join(paragraphs)
    if len(body) > MAX_EXTRACTED_CHARS:
        logger.warning("parse_docx_rejected_text_volume", limit=MAX_EXTRACTED_CHARS)
        raise DocumentParseError(
            "The DOCX expands to more text than this service will process."
        )
    return _assemble([body], DOCX_MIME)


def parse_document(content: bytes, media_type: str) -> ParsedDocument:
    """Parse a resume or job description into text with exact offsets.

    Args:
        content: Raw file bytes.
        media_type: The media type to parse as. Callers should pass the result
            of `app.utils.parsing.detect_media_type`, not a client-supplied
            header.

    Returns:
        The parsed document, with `needs_ocr` set when the text layer is thin.

    Raises:
        EmptyDocumentError: If `content` is empty.
        UnsupportedMediaTypeError: If `media_type` is not PDF or DOCX.
        DocumentParseError: If the document is structurally unreadable.
    """
    if not content:
        raise EmptyDocumentError()
    if media_type == "application/pdf":
        return _parse_pdf(content)
    if media_type == DOCX_MIME:
        return _parse_docx(content)
    raise UnsupportedMediaTypeError()
