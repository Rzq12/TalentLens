"""Shared pytest fixtures for the TalentLens test suite."""

from __future__ import annotations

import os
import uuid
from collections.abc import AsyncIterator, Iterator
from datetime import UTC, datetime, timedelta
from typing import Any

import pytest

JWT_TEST_SECRET = "test-secret-not-a-real-key"
API_KEY_TEST = "test-api-key"


@pytest.fixture(scope="session", autouse=True)
def _test_environment() -> Iterator[None]:
    """Populate the environment `Settings` requires before any import of app.config.

    Values are synthetic and local-only. Nothing here is a real credential.
    """
    defaults = {
        "DATABASE_URL": "postgresql+asyncpg://postgres:postgres@localhost:5433/talentlens_test",
        "JWT_SECRET": JWT_TEST_SECRET,
        "JWT_ISSUER": "https://test.local/auth/v1",
        "JWT_AUDIENCE": "authenticated",
        "STORAGE_BACKEND": "memory",
        "ENVIRONMENT": "test",
    }
    previous = {k: os.environ.get(k) for k in defaults}
    os.environ.update(defaults)
    yield
    for key, value in previous.items():
        if value is None:
            os.environ.pop(key, None)
        else:
            os.environ[key] = value


@pytest.fixture(autouse=True)
def _reset_rate_limiter() -> Iterator[None]:
    """Clear rate-limit state around every test.

    The limiter is process-global by design, so without this one test's
    requests silently consume the next test's budget and unrelated cases fail
    with 429 depending on execution order.
    """
    from app.main import reset_rate_limiter

    reset_rate_limiter()
    yield
    reset_rate_limiter()


@pytest.fixture
def tenant_id() -> uuid.UUID:
    return uuid.UUID("11111111-1111-1111-1111-111111111111")


@pytest.fixture
def other_tenant_id() -> uuid.UUID:
    return uuid.UUID("22222222-2222-2222-2222-222222222222")


@pytest.fixture
def user_id() -> uuid.UUID:
    return uuid.UUID("33333333-3333-3333-3333-333333333333")


@pytest.fixture
def make_token() -> Any:
    """Return a factory that mints signed JWTs for tests.

    The factory accepts overrides so a test can produce an expired token, a
    token signed with the wrong key, or a token missing required claims.
    """
    import jwt

    def _make(
        *,
        sub: str = "33333333-3333-3333-3333-333333333333",
        tenant: str = "11111111-1111-1111-1111-111111111111",
        roles: list[str] | None = None,
        expires_in: timedelta = timedelta(minutes=15),
        secret: str = JWT_TEST_SECRET,
        issuer: str = "https://test.local/auth/v1",
        audience: str = "authenticated",
        drop_claims: tuple[str, ...] = (),
    ) -> str:
        now = datetime.now(UTC)
        payload: dict[str, Any] = {
            "sub": sub,
            "tenant_id": tenant,
            "roles": roles if roles is not None else ["recruiter"],
            "iss": issuer,
            "aud": audience,
            "iat": int(now.timestamp()),
            "exp": int((now + expires_in).timestamp()),
        }
        for claim in drop_claims:
            payload.pop(claim, None)
        return jwt.encode(payload, secret, algorithm="HS256")

    return _make


@pytest.fixture
def auth_headers(make_token: Any) -> dict[str, str]:
    return {"Authorization": f"Bearer {make_token()}"}


@pytest.fixture
async def client() -> AsyncIterator[Any]:
    """An httpx AsyncClient bound to the FastAPI app via ASGI transport."""
    import httpx

    from app.main import create_app

    app = create_app()
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


@pytest.fixture
async def db_client() -> AsyncIterator[Any]:
    """An httpx client backed by a real PostgreSQL schema, created per test.

    The schema is built and dropped around each test so isolation properties
    (tenant scoping, unique constraints) are observed against a clean database
    rather than residue from an earlier case.
    """
    import httpx

    from app.db import Base, dispose_engine, get_engine
    from app.main import create_app

    engine = get_engine()
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.drop_all)
        await connection.run_sync(Base.metadata.create_all)

    app = create_app()
    transport = httpx.ASGITransport(app=app)
    async with httpx.AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac

    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.drop_all)
    await dispose_engine()


class _SchemaInspector:
    """Runs a synchronous callable against a fresh connection.

    `run_sync` lives on `AsyncConnection`, not `AsyncEngine`, so tests that want
    to reflect the schema need a connection opened for them. Wrapping that here
    keeps each test to a single readable call.
    """

    def __init__(self, engine: Any) -> None:
        self._engine = engine

    async def run_sync(self, fn: Any, *args: Any) -> Any:
        """Execute `fn(connection, *args)` on a new connection."""
        async with self._engine.connect() as connection:
            return await connection.run_sync(fn, *args)


@pytest.fixture
async def clean_database() -> AsyncIterator[Any]:
    """Yield a schema inspector over a database with no tables at all.

    Migration tests must start from genuinely empty — including no
    `alembic_version` row — otherwise an upgrade is a no-op and proves nothing.
    """
    from sqlalchemy import text

    from app.db import Base, dispose_engine, get_engine

    async def _wipe() -> None:
        engine = get_engine()
        async with engine.begin() as connection:
            await connection.run_sync(Base.metadata.drop_all)
            await connection.execute(text("DROP TABLE IF EXISTS alembic_version"))

    await _wipe()
    yield _SchemaInspector(get_engine())
    await _wipe()
    await dispose_engine()


@pytest.fixture
def minimal_pdf_bytes() -> bytes:
    """A tiny but structurally valid single-page PDF containing known text."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 720), "Jane Doe")
    page.insert_text((72, 700), "Senior Backend Engineer")
    page.insert_text((72, 680), "Python, Kubernetes, PostgreSQL")
    data: bytes = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def minimal_docx_bytes() -> bytes:
    """A structurally valid DOCX containing known text."""
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("Python, Kubernetes, PostgreSQL")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def scanned_pdf_bytes() -> bytes:
    """A PDF page with no extractable text layer — the OCR-fallback trigger."""
    import fitz

    doc = fitz.open()
    doc.new_page()
    data: bytes = doc.tobytes()
    doc.close()
    return data


@pytest.fixture
def corrupt_pdf_bytes() -> bytes:
    """Bytes carrying a PDF magic number but an unparseable body."""
    return b"%PDF-1.7\n%\xe2\xe3\xcf\xd3\nthis is not a real pdf body"
